import os
import uuid
from PyPDF2 import PdfReader
import docx
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document

persist_directory = "data"

embedder = HuggingFaceEmbeddings(
    model_name = "sentence-transformers/all-MiniLM-L6-v2"
)

vectordb = Chroma(embedding_function=embedder, persist_directory=persist_directory)

def read_file_content(file_path):
    print(f"Reading file: {file_path}")
    if not os.path.exists(file_path):
        raise FileNotFoundError("File does not exist.")
    
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        try:
            with open(file_path, 'r') as file:
                content = file.read()
                print('File content:', content)
                return content
        except Exception as e:
            print(f"Error reading .txt file: {e}")
            raise e
    elif ext == ".pdf":
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                content = ""
                for page in reader.pages:
                    page_content = page.extract_text()
                    print(f"Page content: {page_content}")
                    content += page_content
                return content
        except Exception as e:
            print(f"Error reading .pdf file: {e}")
            raise e
    elif ext == ".docx":
        try:
            print(f"Attempting to read .docx file: {file_path}")
            doc = docx.Document(file_path)
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            content = '\n'.join(fullText)
            print('File content:', content)
            return content
        except Exception as e:
            print(f"Error reading .docx file: {e}")
            raise e
    else:
        raise ValueError("Unsupported file type")

def add_document_to_db(content, metadata):
    asset_id = str(uuid.uuid4())
    documents = [Document(page_content=content, metadata=metadata, id=asset_id)]
    vectordb.add_documents(documents=documents)
    return asset_id

def get_document_from_db(asset_id):
    results = vectordb.get(ids=[asset_id])
    return results

def process_document_service(file_path: str):
    content = read_file_content(file_path)
    metadata = {"file_path": file_path}
    asset_id = add_document_to_db(content, metadata)
    get_document_from_db(asset_id)
    return asset_id